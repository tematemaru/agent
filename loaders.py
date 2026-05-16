import os
import json
import pandas as pd

from bs4 import BeautifulSoup

from docx import Document

from langchain_core.documents import Document as LCDocument

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader
)


TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".json",
    ".yaml",
    ".yml",
    ".xml",
    ".html",
    ".css",
    ".sql",
    ".sh"
}


def load_txt(path):

    loader = TextLoader(
        path,
        encoding="utf-8"
    )

    return loader.load()


def load_pdf(path):

    loader = PyPDFLoader(path)

    return loader.load()


def load_docx(path):

    doc = Document(path)

    text = "\n".join(
        [p.text for p in doc.paragraphs]
    )

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": os.path.basename(path)
            }
        )
    ]


def load_csv(path):

    df = pd.read_csv(path)

    text = df.to_csv(index=False)

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": os.path.basename(path)
            }
        )
    ]


def load_xlsx(path):

    excel = pd.ExcelFile(path)

    chunks = []

    for sheet in excel.sheet_names:

        df = excel.parse(sheet)

        text = df.to_csv(index=False)

        chunks.append(
            LCDocument(
                page_content=text,
                metadata={
                    "source": os.path.basename(path),
                    "sheet": sheet
                }
            )
        )

    return chunks


def load_json(path):

    with open(path, "r", encoding="utf-8") as f:

        data = json.load(f)

    text = json.dumps(
        data,
        ensure_ascii=False,
        indent=2
    )

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": os.path.basename(path)
            }
        )
    ]


def load_html(path):

    with open(path, "r", encoding="utf-8") as f:

        html = f.read()

    soup = BeautifulSoup(
        html,
        "lxml"
    )

    text = soup.get_text("\n")

    return [
        LCDocument(
            page_content=text,
            metadata={
                "source": os.path.basename(path)
            }
        )
    ]


def load_file(path):

    ext = os.path.splitext(path)[1].lower()

    try:

        if ext in TEXT_EXTENSIONS:
            return load_txt(path)

        if ext == ".pdf":
            return load_pdf(path)

        if ext == ".docx":
            return load_docx(path)

        if ext == ".csv":
            return load_csv(path)

        if ext == ".xlsx":
            return load_xlsx(path)

        if ext == ".json":
            return load_json(path)

        if ext in [".html", ".htm"]:
            return load_html(path)

        return []

    except Exception as e:

        print(f"❌ failed loading {path}: {e}")

        return []